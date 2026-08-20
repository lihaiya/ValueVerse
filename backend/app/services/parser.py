from dataclasses import dataclass
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class ParsedDocument:
    text: str
    parser_name: str
    warnings: list[str]
    spans: list[dict[str, Any]]
    quality: dict[str, Any]


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported file type: {suffix}")
    if suffix in {".txt", ".md"}:
        text = _read_text(path)
        spans = _text_spans(text)
        return ParsedDocument(
            text=text,
            parser_name="plain-text",
            warnings=[],
            spans=spans,
            quality=_build_quality(text=text, warnings=[], spans=spans, parser_name="plain-text"),
        )
    if suffix == ".docx":
        return _parse_docx(path)
    return _parse_pdf(path)


def _read_text(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to parse DOCX files") from exc
    document = Document(str(path))
    spans: list[dict[str, Any]] = []
    paragraphs: list[str] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        spans.append({"type": "paragraph", "locator": {"paragraph": index + 1}, "text": text, "confidence": 1.0})
    table_lines: list[str] = []
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            text = " | ".join(cell.text.strip() for cell in row.cells)
            if not text.strip():
                continue
            table_lines.append(text)
            spans.append(
                {
                    "type": "table_row",
                    "locator": {"table": table_index + 1, "row": row_index + 1},
                    "text": text,
                    "confidence": 1.0,
                }
            )
    full_text = "\n\n".join(paragraphs + table_lines)
    return ParsedDocument(
        text=full_text,
        parser_name="python-docx",
        warnings=[],
        spans=spans,
        quality=_build_quality(text=full_text, warnings=[], spans=spans, parser_name="python-docx"),
    )


def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pypdf is required to parse PDF files") from exc
    reader = PdfReader(str(path))
    warnings: list[str] = []
    pages: list[str] = []
    spans: list[dict[str, Any]] = []
    for index, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover
            warnings.append(f"page {index + 1}: {exc}")
            text = ""
        if text.strip():
            clean_text = text.strip()
            pages.append(f"<!-- page:{index + 1} -->\n{clean_text}")
            spans.append({"type": "page", "locator": {"page": index + 1}, "text": clean_text, "confidence": 0.85})
    full_text = "\n\n".join(pages)
    return ParsedDocument(
        text=full_text,
        parser_name="pypdf",
        warnings=warnings,
        spans=spans,
        quality=_build_quality(
            text=full_text,
            warnings=warnings,
            spans=spans,
            parser_name="pypdf",
            page_count=len(reader.pages),
        ),
    )


def _text_spans(text: str) -> list[dict[str, Any]]:
    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n") if paragraph.strip()]
    if not paragraphs and text.strip():
        paragraphs = [text.strip()]
    return [
        {"type": "paragraph", "locator": {"paragraph": index + 1}, "text": paragraph, "confidence": 1.0}
        for index, paragraph in enumerate(paragraphs[:500])
    ]


def _build_quality(
    *,
    text: str,
    warnings: list[str],
    spans: list[dict[str, Any]],
    parser_name: str,
    page_count: int | None = None,
) -> dict[str, Any]:
    text_length = len(text.strip())
    replacement_ratio = text.count("\ufffd") / max(1, len(text))
    average_chars_per_page = text_length / page_count if page_count else None
    needs_ocr = parser_name == "pypdf" and (
        text_length == 0
        or bool(page_count and average_chars_per_page is not None and average_chars_per_page < 120)
        or bool(page_count and len(spans) / max(1, page_count) < 0.5)
    )
    status = "ok"
    if warnings or replacement_ratio > 0.01 or needs_ocr or text_length == 0:
        status = "warning"
    return {
        "status": status,
        "text_length": text_length,
        "span_count": len(spans),
        "page_count": page_count,
        "average_chars_per_page": average_chars_per_page,
        "replacement_ratio": replacement_ratio,
        "needs_ocr": needs_ocr,
        "ocr_candidates": ["hunyuanocr"] if needs_ocr else [],
        "warnings": warnings,
    }
